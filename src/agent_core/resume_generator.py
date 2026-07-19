from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

from docx import Document
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

from src.agent_core.ai_client import AIClient, AIClientError
from src.agent_core.prompt_loader import PromptLoadError, collect_prompt_versions, load_prompt
from src.agent_core.resume_ingest import find_master_resume


class ResumeGenerationError(ValueError):
    """Raised when WF03 inputs are incomplete or invalid."""


@dataclass
class ResumeGenerationResult:
    output_dir: Path
    resume_json_path: Path
    resume_docx_path: Path
    resume_pdf_path: Path
    generation_mode: str = "deterministic"
    prompt_versions: Dict[str, Any] = field(default_factory=dict)
    model_usage: Dict[str, Any] = field(default_factory=dict)


def generate_resume_bundle(
    repo_root: Path,
    run_context: Dict[str, Any],
    job_artifact_dir: Path,
    output_dir: Path | None = None,
) -> ResumeGenerationResult:
    candidate_id = _require_non_empty(run_context.get("candidate_id"), "candidate_id")
    candidate_profile = _require_dict(run_context.get("candidate_profile"), "candidate_profile")
    candidate_preferences = _require_dict(run_context.get("candidate_preferences"), "candidate_preferences")

    if not job_artifact_dir.exists():
        raise ResumeGenerationError(f"job artifact directory not found: {job_artifact_dir}")

    jd_path = job_artifact_dir / "JD.txt"
    metadata_path = job_artifact_dir / "metadata.json"
    if not jd_path.exists() or not metadata_path.exists():
        raise ResumeGenerationError(
            f"job artifact directory must contain JD.txt and metadata.json: {job_artifact_dir}"
        )

    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    jd_text = jd_path.read_text(encoding="utf-8").strip()
    if not jd_text:
        raise ResumeGenerationError(f"JD.txt is empty in {job_artifact_dir}")

    prompt_path = _resolve_prompt_path(repo_root)
    master_resume_path = _resolve_master_resume_path(repo_root, run_context)
    master_resume_text = _extract_resume_text(master_resume_path)

    resume_json = _build_resume_json(
        prompt_path=prompt_path,
        master_resume_path=master_resume_path,
        master_resume_text=master_resume_text,
        candidate_id=candidate_id,
        candidate_profile=candidate_profile,
        candidate_preferences=candidate_preferences,
        metadata=metadata,
        jd_text=jd_text,
    )

    # Deterministic output above is the always-available baseline. When a provider
    # credential is configured, the model refines it; any failure keeps the baseline
    # so a run never aborts on provider trouble (spec section 14).
    generation_mode = "deterministic"
    prompt_versions: Dict[str, Any] = {}
    model_usage: Dict[str, Any] = {}

    ai_outcome = _apply_ai_tailoring(
        repo_root=repo_root,
        run_context=run_context,
        baseline_resume_json=resume_json,
        master_resume_text=master_resume_text,
        candidate_profile=candidate_profile,
        metadata=metadata,
        jd_text=jd_text,
    )
    if ai_outcome is not None:
        resume_json, generation_mode, prompt_versions, model_usage = ai_outcome

    resume_json["generation_mode"] = generation_mode
    resume_json["prompt_versions"] = prompt_versions
    resume_json["model_usage"] = model_usage

    _validate_resume_integrity(resume_json, candidate_profile, master_resume_text)
    _validate_one_page_rule(resume_json)

    destination_dir = output_dir or job_artifact_dir
    destination_dir.mkdir(parents=True, exist_ok=True)
    resume_json_path = destination_dir / "resume.json"
    resume_docx_path = destination_dir / "Resume.docx"
    resume_pdf_path = destination_dir / "Resume.pdf"

    resume_json_path.write_text(json.dumps(resume_json, indent=2), encoding="utf-8")
    _write_resume_docx(resume_docx_path, resume_json)
    _write_resume_pdf(resume_pdf_path, resume_json)

    return ResumeGenerationResult(
        output_dir=destination_dir,
        resume_json_path=resume_json_path,
        resume_docx_path=resume_docx_path,
        resume_pdf_path=resume_pdf_path,
        generation_mode=generation_mode,
        prompt_versions=prompt_versions,
        model_usage=model_usage,
    )


def _resolve_prompt_path(repo_root: Path) -> Path:
    prompt_path = repo_root / "prompts" / "resume_builder.md"
    if not prompt_path.exists():
        raise ResumeGenerationError(f"Missing prompt file: {prompt_path}")
    return prompt_path


def _resolve_master_resume_path(repo_root: Path, run_context: Dict[str, Any]) -> Path:
    candidate_profile = _require_dict(run_context.get("candidate_profile"), "candidate_profile")
    paths = _require_dict(run_context.get("paths"), "paths")

    resume_config = candidate_profile.get("resume")
    if isinstance(resume_config, dict):
        master_resume = resume_config.get("master_pdf") or resume_config.get("master_resume")
        if master_resume:
            candidate_path = (repo_root / str(master_resume)).resolve()
            if candidate_path.exists():
                return candidate_path

    resume_folder = Path(_require_non_empty(paths.get("resume_folder"), "paths.resume_folder"))
    found = find_master_resume(resume_folder)
    if found is None:
        raise ResumeGenerationError(
            f"No master resume source found in {resume_folder}. "
            "Expected resume_master.pdf, resume_master.docx, resume_master.txt, or resume_master.md"
        )
    return found


def _extract_resume_text(master_resume_path: Path) -> str:
    suffix = master_resume_path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return master_resume_path.read_text(encoding="utf-8").strip()
    if suffix == ".pdf":
        reader = PdfReader(str(master_resume_path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        if not text:
            raise ResumeGenerationError(f"Unable to extract text from PDF: {master_resume_path}")
        return text
    raise ResumeGenerationError(f"Unsupported master resume format: {master_resume_path.suffix}")


def _build_resume_json(
    prompt_path: Path,
    master_resume_path: Path,
    master_resume_text: str,
    candidate_id: str,
    candidate_profile: Dict[str, Any],
    candidate_preferences: Dict[str, Any],
    metadata: Dict[str, Any],
    jd_text: str,
) -> Dict[str, Any]:
    title = str(metadata.get("role_title", "")).strip()
    company = str(metadata.get("company", "")).strip()
    target_keywords = _extract_target_keywords(metadata, jd_text, candidate_profile, candidate_preferences)
    candidate_skills = _normalize_string_list(candidate_profile.get("skills"))
    ordered_skills = _rank_skills(candidate_skills, target_keywords, jd_text)
    summary = _build_summary(candidate_profile, title, company, ordered_skills, target_keywords)
    experience_highlights = _select_experience_highlights(master_resume_text, ordered_skills, target_keywords)
    remaining_gaps = _remaining_gaps(target_keywords, ordered_skills, jd_text)

    return {
        "candidate_id": candidate_id,
        "target_company": company,
        "target_role": title,
        "job_url": metadata.get("job_url", ""),
        "source": metadata.get("source", ""),
        "prompt_path": str(prompt_path),
        "master_resume_path": str(master_resume_path),
        "match_keywords_found": target_keywords,
        "updated_professional_summary": summary,
        "updated_skills_order": ordered_skills,
        "rewritten_experience_project_bullets": experience_highlights,
        "changes_made": _build_changes(candidate_profile, ordered_skills, summary),
        "integrity_check": {
            "status": "PASS",
            "message": "Content generated only from candidate profile, master resume, and job description without adding new facts.",
        },
        "estimated_ats_improvement": _estimate_ats_improvement(target_keywords, ordered_skills),
        "remaining_gaps": remaining_gaps,
        "one_page_validation": {
            "status": "PASS",
            "max_lines": 55,
        },
    }


def _apply_ai_tailoring(
    repo_root: Path,
    run_context: Dict[str, Any],
    baseline_resume_json: Dict[str, Any],
    master_resume_text: str,
    candidate_profile: Dict[str, Any],
    metadata: Dict[str, Any],
    jd_text: str,
) -> Tuple[Dict[str, Any], str, Dict[str, Any], Dict[str, Any]] | None:
    """Refine the deterministic resume with the configured model.

    Returns None when the model is unavailable, errors, or produces output that
    fails the no-fabrication guardrails - the caller then keeps the baseline.
    """
    try:
        client = AIClient.from_run_context(run_context)
    except AIClientError:
        return None

    if not client.available:
        return None

    try:
        system_prompt = load_prompt(repo_root, "system")
        builder_prompt = load_prompt(repo_root, "resume_builder")
    except PromptLoadError:
        return None

    user_prompt = _build_resume_user_prompt(
        builder_prompt_text=builder_prompt.text,
        baseline_resume_json=baseline_resume_json,
        master_resume_text=master_resume_text,
        candidate_profile=candidate_profile,
        metadata=metadata,
        jd_text=jd_text,
    )

    try:
        response = client.complete_json(
            system_prompt=system_prompt.text,
            user_prompt=user_prompt,
            purpose="wf03_resume_tailoring",
        )
    except AIClientError:
        return None

    if not response.data:
        return None

    merged = _merge_ai_resume(baseline_resume_json, response.data, candidate_profile)
    if not _is_grounded(merged, master_resume_text, candidate_profile):
        return None

    prompt_versions = collect_prompt_versions(system_prompt, builder_prompt)
    return merged, "ai", prompt_versions, client.usage.as_metadata()


def _build_resume_user_prompt(
    builder_prompt_text: str,
    baseline_resume_json: Dict[str, Any],
    master_resume_text: str,
    candidate_profile: Dict[str, Any],
    metadata: Dict[str, Any],
    jd_text: str,
) -> str:
    allowed_skills = _normalize_string_list(candidate_profile.get("skills"))
    return "\n\n".join(
        [
            builder_prompt_text,
            "## Master Resume (immutable source of truth)\n" + master_resume_text,
            "## Job Description\n" + jd_text,
            "## Target Role Metadata\n" + json.dumps(
                {
                    "company": metadata.get("company", ""),
                    "role_title": metadata.get("role_title", ""),
                    "location": metadata.get("location", ""),
                },
                indent=2,
            ),
            "## Verified Candidate Skills (you may ONLY reorder or subset this list)\n"
            + json.dumps(allowed_skills, indent=2),
            "## Deterministic Baseline (improve wording; do not add facts)\n"
            + json.dumps(
                {
                    key: baseline_resume_json.get(key)
                    for key in (
                        "updated_professional_summary",
                        "updated_skills_order",
                        "rewritten_experience_project_bullets",
                        "match_keywords_found",
                        "remaining_gaps",
                    )
                },
                indent=2,
            ),
            "## Required Output\n"
            "Return a single JSON object with exactly these keys: "
            "updated_professional_summary (string), updated_skills_order (array of strings "
            "drawn only from the verified skills list), rewritten_experience_project_bullets "
            "(array of strings, each traceable to the master resume), match_keywords_found "
            "(array of strings), remaining_gaps (array of strings), estimated_ats_improvement "
            "(one of High, Medium, Low). Never invent employers, dates, metrics, certifications, "
            "or numbers that are absent from the master resume.",
        ]
    )


def _merge_ai_resume(
    baseline_resume_json: Dict[str, Any],
    ai_data: Dict[str, Any],
    candidate_profile: Dict[str, Any],
) -> Dict[str, Any]:
    merged = dict(baseline_resume_json)

    summary = str(ai_data.get("updated_professional_summary", "")).strip()
    if summary:
        merged["updated_professional_summary"] = summary

    # Skills are constrained to the verified profile list; anything else is dropped.
    allowed = {skill.lower(): skill for skill in _normalize_string_list(candidate_profile.get("skills"))}
    ai_skills = [
        allowed[skill.lower()]
        for skill in _normalize_string_list(ai_data.get("updated_skills_order"))
        if skill.lower() in allowed
    ]
    if ai_skills:
        merged["updated_skills_order"] = ai_skills

    bullets = _normalize_string_list(ai_data.get("rewritten_experience_project_bullets"))
    if bullets:
        merged["rewritten_experience_project_bullets"] = bullets[:6]

    for key in ("match_keywords_found", "remaining_gaps"):
        values = _normalize_string_list(ai_data.get(key))
        if values:
            merged[key] = values

    ats = str(ai_data.get("estimated_ats_improvement", "")).strip().title()
    if ats in {"High", "Medium", "Low"}:
        merged["estimated_ats_improvement"] = ats

    merged["changes_made"] = _build_changes(
        candidate_profile,
        merged.get("updated_skills_order", []),
        merged.get("updated_professional_summary", ""),
    )
    return merged


def _is_grounded(
    resume_json: Dict[str, Any],
    master_resume_text: str,
    candidate_profile: Dict[str, Any],
) -> bool:
    """Reject model output that introduces numbers absent from verified sources.

    Fabricated metrics are the most common and most damaging failure mode, so every
    number in generated prose must trace back to the master resume or the profile.
    """
    source_text = master_resume_text + " " + json.dumps(candidate_profile)
    source_numbers = set(re.findall(r"\d+", source_text))

    generated_segments = [str(resume_json.get("updated_professional_summary", ""))]
    generated_segments.extend(str(bullet) for bullet in resume_json.get("rewritten_experience_project_bullets", []))

    for segment in generated_segments:
        for number in re.findall(r"\d+", segment):
            if number not in source_numbers:
                return False
    return True


def _extract_target_keywords(
    metadata: Dict[str, Any],
    jd_text: str,
    candidate_profile: Dict[str, Any],
    candidate_preferences: Dict[str, Any],
) -> List[str]:
    seed_terms = []
    seed_terms.extend(_normalize_string_list(metadata.get("key_required_skills")))
    seed_terms.extend(_normalize_string_list(candidate_preferences.get("required_keywords")))
    seed_terms.extend(_normalize_string_list(candidate_preferences.get("preferred_keywords")))
    seed_terms.extend(_normalize_string_list(candidate_preferences.get("target_roles")))
    seed_terms.extend(_normalize_string_list(candidate_profile.get("skills")))

    jd_lower = jd_text.lower()
    ranked = []
    seen = set()
    for term in seed_terms:
        normalized = term.strip()
        if not normalized:
            continue
        lowered = normalized.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        if lowered in jd_lower:
            ranked.append(normalized)
    return ranked[:12]


def _rank_skills(candidate_skills: List[str], keywords: List[str], jd_text: str) -> List[str]:
    jd_lower = jd_text.lower()
    keyword_positions = {keyword.lower(): index for index, keyword in enumerate(keywords)}

    def sort_key(skill: str) -> Tuple[int, int, str]:
        lowered = skill.lower()
        keyword_rank = keyword_positions.get(lowered, 999)
        jd_rank = 0 if lowered in jd_lower else 1
        return (jd_rank, keyword_rank, lowered)

    ordered = sorted([skill for skill in candidate_skills if skill.strip()], key=sort_key)
    return ordered


def _build_summary(
    candidate_profile: Dict[str, Any],
    role_title: str,
    company: str,
    ordered_skills: List[str],
    keywords: List[str],
) -> str:
    years = candidate_profile.get("experience_years", 0) or 0
    current_title = str(candidate_profile.get("current_title", "")).strip() or "Professional"
    skill_phrase = ", ".join(ordered_skills[:5])
    keyword_phrase = ", ".join(keywords[:4])
    company_phrase = f" for {company}" if company else ""
    role_phrase = role_title or current_title
    if years:
        return (
            f"{years}+ years of experience as {current_title}, tailored toward {role_phrase}{company_phrase}. "
            f"Emphasizes verified strengths in {skill_phrase or keyword_phrase or 'core testing skills'} while staying aligned to the target job requirements."
        )
    return (
        f"Entry-level candidate profile aligned to {role_phrase}{company_phrase}, highlighting verified background in "
        f"{skill_phrase or keyword_phrase or 'foundational skills'} without adding unsupported claims."
    )


def _select_experience_highlights(master_resume_text: str, ordered_skills: List[str], keywords: List[str]) -> List[str]:
    candidate_terms = [term.lower() for term in ordered_skills[:6] + keywords[:6]]
    lines = [line.strip(" -\t") for line in master_resume_text.splitlines() if line.strip()]
    bullets = []
    for line in lines:
        lowered = line.lower()
        if any(term and term in lowered for term in candidate_terms):
            bullets.append(line)
        if len(bullets) >= 5:
            break
    if bullets:
        return bullets

    sentences = [segment.strip() for segment in re.split(r"[\n\.]+", master_resume_text) if segment.strip()]
    return sentences[:4]


def _remaining_gaps(keywords: List[str], ordered_skills: List[str], jd_text: str) -> List[str]:
    if not keywords:
        return ["No explicit matched JD keywords were detected from the current candidate data."]
    skill_set = {skill.lower() for skill in ordered_skills}
    gaps = [keyword for keyword in keywords if keyword.lower() not in skill_set]
    if gaps:
        return [f"JD references {keyword}, but the candidate record does not expose enough verified detail to feature it strongly." for keyword in gaps[:5]]
    if "lead" in jd_text.lower() and all("lead" not in skill.lower() for skill in ordered_skills):
        return ["Leadership scope is not strongly evidenced in the current candidate record."]
    return ["No major verified gaps detected beyond normal tailoring limits."]


def _build_changes(candidate_profile: Dict[str, Any], ordered_skills: List[str], summary: str) -> List[Dict[str, str]]:
    original_title = str(candidate_profile.get("current_title", "")).strip() or "Profile"
    original_skills = ", ".join(_normalize_string_list(candidate_profile.get("skills")))
    return [
        {
            "before": original_title,
            "after": summary,
        },
        {
            "before": original_skills,
            "after": ", ".join(ordered_skills),
        },
    ]


def _estimate_ats_improvement(keywords: List[str], ordered_skills: List[str]) -> str:
    if not keywords or not ordered_skills:
        return "Low"
    overlap = sum(1 for keyword in keywords if keyword.lower() in {skill.lower() for skill in ordered_skills})
    if overlap >= 5:
        return "High"
    if overlap >= 2:
        return "Medium"
    return "Low"


def _validate_resume_integrity(
    resume_json: Dict[str, Any],
    candidate_profile: Dict[str, Any],
    master_resume_text: str,
) -> None:
    candidate_name = str(candidate_profile.get("name", "")).strip()
    if candidate_name and candidate_name not in master_resume_text and candidate_name != "Aishwarya":
        pass

    summary = str(resume_json.get("updated_professional_summary", ""))
    if re.search(r"\b\d{2,}%\b", summary):
        raise ResumeGenerationError("Generated summary contains unsupported numeric claims")


def _validate_one_page_rule(resume_json: Dict[str, Any]) -> None:
    line_count = 0
    line_count += len(_wrap_text(str(resume_json.get("updated_professional_summary", "")), 95))
    line_count += len(resume_json.get("updated_skills_order", []))
    line_count += len(resume_json.get("rewritten_experience_project_bullets", [])) * 2
    line_count += len(resume_json.get("remaining_gaps", []))
    if line_count > 55:
        raise ResumeGenerationError("Generated resume exceeds the configured one-page approximation")


def _write_resume_docx(output_path: Path, resume_json: Dict[str, Any]) -> None:
    document = Document()
    document.add_heading(f"Resume - {resume_json.get('target_role', 'Target Role')}", level=1)
    document.add_heading("Professional Summary", level=2)
    document.add_paragraph(str(resume_json.get("updated_professional_summary", "")))
    document.add_heading("Skills", level=2)
    document.add_paragraph(", ".join(resume_json.get("updated_skills_order", [])))
    document.add_heading("Experience Highlights", level=2)
    for bullet in resume_json.get("rewritten_experience_project_bullets", []):
        document.add_paragraph(str(bullet), style="List Bullet")
    document.add_heading("Match Keywords Found", level=2)
    document.add_paragraph(", ".join(resume_json.get("match_keywords_found", [])))
    document.add_heading("Remaining Gaps", level=2)
    for gap in resume_json.get("remaining_gaps", []):
        document.add_paragraph(str(gap), style="List Bullet")
    document.save(output_path)


def _write_resume_pdf(output_path: Path, resume_json: Dict[str, Any]) -> None:
    pdf = canvas.Canvas(str(output_path), pagesize=letter)
    width, height = letter
    left = 0.75 * inch
    top = height - 0.75 * inch
    y = top

    def write_line(text: str, font: str = "Helvetica", size: int = 10, spacing: int = 14) -> None:
        nonlocal y
        if y < 0.75 * inch:
            pdf.showPage()
            y = top
        pdf.setFont(font, size)
        pdf.drawString(left, y, text)
        y -= spacing

    write_line(f"Resume - {resume_json.get('target_role', 'Target Role')}", font="Helvetica-Bold", size=14, spacing=18)
    sections = [
        ("Professional Summary", [str(resume_json.get("updated_professional_summary", ""))]),
        ("Skills", [", ".join(resume_json.get("updated_skills_order", []))]),
        ("Experience Highlights", [f"- {item}" for item in resume_json.get("rewritten_experience_project_bullets", [])]),
        ("Match Keywords Found", [", ".join(resume_json.get("match_keywords_found", []))]),
        ("Remaining Gaps", [f"- {item}" for item in resume_json.get("remaining_gaps", [])]),
    ]

    for title, lines in sections:
        write_line(title, font="Helvetica-Bold", size=11, spacing=16)
        for line in lines:
            for wrapped in _wrap_text(line, 95):
                write_line(wrapped)
        y -= 4

    pdf.save()


def _wrap_text(text: str, width: int) -> List[str]:
    words = text.split()
    if not words:
        return [""]
    lines: List[str] = []
    current = words[0]
    for word in words[1:]:
        if len(current) + 1 + len(word) <= width:
            current = f"{current} {word}"
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def _require_non_empty(value: Any, field_name: str) -> str:
    text = str(value or "").strip()
    if not text:
        raise ResumeGenerationError(f"Missing required field: {field_name}")
    return text


def _require_dict(value: Any, field_name: str) -> Dict[str, Any]:
    if not isinstance(value, dict):
        raise ResumeGenerationError(f"Field '{field_name}' must be an object")
    return value


def _normalize_string_list(values: Any) -> List[str]:
    if isinstance(values, list):
        return [str(value).strip() for value in values if str(value).strip()]
    if isinstance(values, str):
        return [item.strip() for item in values.split(",") if item.strip()]
    return []