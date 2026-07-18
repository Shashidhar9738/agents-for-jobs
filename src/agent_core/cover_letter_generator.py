from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

from src.agent_core.ai_client import AIClient, AIClientError
from src.agent_core.prompt_loader import PromptLoadError, collect_prompt_versions, load_prompt

MIN_WORDS = 220
MAX_WORDS = 320


class CoverLetterGenerationError(ValueError):
    """Raised when WF04 inputs are incomplete or invalid."""


@dataclass
class CoverLetterGenerationResult:
    output_dir: Path
    cover_letter_text_path: Path
    cover_letter_docx_path: Path
    cover_letter_pdf_path: Path
    generation_mode: str = "deterministic"
    prompt_versions: Dict[str, Any] = field(default_factory=dict)
    model_usage: Dict[str, Any] = field(default_factory=dict)


def generate_cover_letter_bundle(
    repo_root: Path,
    run_context: Dict[str, Any],
    job_artifact_dir: Path,
    output_dir: Path | None = None,
) -> CoverLetterGenerationResult:
    candidate_id = _require_text(run_context.get("candidate_id"), "candidate_id")
    candidate_profile = _require_dict(run_context.get("candidate_profile"), "candidate_profile")

    prompt_path = repo_root / "prompts" / "cover_letter.md"
    if not prompt_path.exists():
        raise CoverLetterGenerationError(f"Missing prompt file: {prompt_path}")

    jd_path = job_artifact_dir / "JD.txt"
    metadata_path = job_artifact_dir / "metadata.json"
    resume_json_path = job_artifact_dir / "resume.json"
    if not jd_path.exists() or not metadata_path.exists() or not resume_json_path.exists():
        raise CoverLetterGenerationError(
            f"WF04 requires JD.txt, metadata.json, and resume.json in {job_artifact_dir}"
        )

    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    resume_json = json.loads(resume_json_path.read_text(encoding="utf-8"))
    jd_text = jd_path.read_text(encoding="utf-8").strip()
    cover_letter_text = _build_cover_letter(candidate_profile, metadata, resume_json, jd_text)
    _validate_cover_letter(cover_letter_text, candidate_profile)

    # Deterministic letter above is the guaranteed-valid baseline; the model only
    # replaces it if its output also passes every validation rule.
    generation_mode = "deterministic"
    prompt_versions: Dict[str, Any] = {}
    model_usage: Dict[str, Any] = {}

    ai_outcome = _apply_ai_cover_letter(
        repo_root=repo_root,
        run_context=run_context,
        candidate_profile=candidate_profile,
        metadata=metadata,
        resume_json=resume_json,
        jd_text=jd_text,
    )
    if ai_outcome is not None:
        cover_letter_text, generation_mode, prompt_versions, model_usage = ai_outcome

    destination_dir = output_dir or job_artifact_dir
    destination_dir.mkdir(parents=True, exist_ok=True)
    text_path = destination_dir / "CoverLetter.txt"
    docx_path = destination_dir / "CoverLetter.docx"
    pdf_path = destination_dir / "CoverLetter.pdf"

    text_path.write_text(cover_letter_text, encoding="utf-8")
    _write_docx(docx_path, cover_letter_text, metadata)
    _write_pdf(pdf_path, cover_letter_text, metadata)

    return CoverLetterGenerationResult(
        output_dir=destination_dir,
        cover_letter_text_path=text_path,
        cover_letter_docx_path=docx_path,
        cover_letter_pdf_path=pdf_path,
        generation_mode=generation_mode,
        prompt_versions=prompt_versions,
        model_usage=model_usage,
    )


def _apply_ai_cover_letter(
    repo_root: Path,
    run_context: Dict[str, Any],
    candidate_profile: Dict[str, Any],
    metadata: Dict[str, Any],
    resume_json: Dict[str, Any],
    jd_text: str,
) -> tuple[str, str, Dict[str, Any], Dict[str, Any]] | None:
    """Generate the letter with the configured model.

    Returns None whenever the model is unavailable, fails, or produces a letter
    that breaks the word-count or no-fabrication rules.
    """
    try:
        client = AIClient.from_run_context(run_context)
    except AIClientError:
        return None

    if not client.available:
        return None

    try:
        system_prompt = load_prompt(repo_root, "system")
        letter_prompt = load_prompt(repo_root, "cover_letter")
    except PromptLoadError:
        return None

    user_prompt = "\n\n".join(
        [
            letter_prompt.text,
            "## Job Description\n" + jd_text,
            "## Target Role Metadata\n" + json.dumps(
                {
                    "company": metadata.get("company", ""),
                    "role_title": metadata.get("role_title", ""),
                    "location": metadata.get("location", ""),
                    "source": metadata.get("source", ""),
                },
                indent=2,
            ),
            "## Candidate Profile (verified facts - do not exceed these)\n"
            + json.dumps(candidate_profile, indent=2),
            "## Tailored Resume Summary\n" + json.dumps(
                {
                    key: resume_json.get(key)
                    for key in (
                        "updated_professional_summary",
                        "updated_skills_order",
                        "match_keywords_found",
                    )
                },
                indent=2,
            ),
            "## Required Output\n"
            f"Return a JSON object with a single key 'cover_letter_text' containing the full letter "
            f"as a string. The letter must be between {MIN_WORDS} and {MAX_WORDS} words, must open with "
            f"'Dear Hiring Team,', must close with the candidate's exact name "
            f"'{candidate_profile.get('name', '')}', and must contain no percentages, invented metrics, "
            f"employers, dates, or certifications absent from the candidate profile and resume summary.",
        ]
    )

    try:
        response = client.complete_json(
            system_prompt=system_prompt.text,
            user_prompt=user_prompt,
            purpose="wf04_cover_letter",
        )
    except AIClientError:
        return None

    if not response.data:
        return None

    letter_text = str(response.data.get("cover_letter_text", "")).strip()
    if not letter_text:
        return None

    if not _is_letter_grounded(letter_text, candidate_profile, resume_json, jd_text):
        return None

    try:
        _validate_cover_letter(letter_text, candidate_profile)
    except CoverLetterGenerationError:
        return None

    prompt_versions = collect_prompt_versions(system_prompt, letter_prompt)
    return letter_text, "ai", prompt_versions, client.usage.as_metadata()


def _is_letter_grounded(
    letter_text: str,
    candidate_profile: Dict[str, Any],
    resume_json: Dict[str, Any],
    jd_text: str,
) -> bool:
    """Reject letters introducing numbers absent from the verified source material."""
    source_text = " ".join(
        [json.dumps(candidate_profile), json.dumps(resume_json), jd_text]
    )
    source_numbers = set(re.findall(r"\d+", source_text))
    return all(number in source_numbers for number in re.findall(r"\d+", letter_text))


def _build_cover_letter(
    candidate_profile: Dict[str, Any],
    metadata: Dict[str, Any],
    resume_json: Dict[str, Any],
    jd_text: str,
) -> str:
    candidate_name = str(candidate_profile.get("name", "Candidate")).strip() or "Candidate"
    current_title = str(candidate_profile.get("current_title", "Professional")).strip() or "Professional"
    company = str(metadata.get("company", "your team")).strip() or "your team"
    role = str(metadata.get("role_title", "the role")).strip() or "the role"
    location = str(metadata.get("location", "")).strip()
    matched_keywords = resume_json.get("match_keywords_found", [])[:4]
    skills = resume_json.get("updated_skills_order", [])[:5]
    summary = str(resume_json.get("updated_professional_summary", "")).strip()
    bullet = ", ".join(skills or matched_keywords or ["relevant testing skills"])
    keyword_phrase = ", ".join(matched_keywords) if matched_keywords else bullet
    location_sentence = ""
    if location:
        location_sentence = f" I am also aligned with the role's location and work-mode expectations in {location}."

    paragraphs = [
        (
            f"Dear Hiring Team,\n\nI am writing to express interest in the {role} opportunity at {company}. "
            f"The role stands out because it emphasizes {keyword_phrase}, which matches the verified experience I have built as a {current_title}."
        ),
        (
            f"Across my background, I have focused on {bullet}. {summary} "
            f"That evidence-backed alignment is the main reason I believe I can contribute effectively in this position."
        ),
        (
            f"What makes this opportunity especially relevant is the overlap between the job description and the work I can support today: {keyword_phrase}."
            f" I would bring a practical, quality-focused approach grounded in the experience already reflected in my resume.{location_sentence}"
        ),
        (
            f"Thank you for your time and consideration. I would welcome the opportunity to discuss how my background can support {company}'s goals in the {role} position.\n\n"
            f"Sincerely,\n{candidate_name}"
        ),
    ]

    text = "\n\n".join(paragraphs)
    return _fit_word_window(text, MIN_WORDS, MAX_WORDS, candidate_profile, metadata, resume_json)


def _fit_word_window(
    text: str,
    min_words: int,
    max_words: int,
    candidate_profile: Dict[str, Any],
    metadata: Dict[str, Any],
    resume_json: Dict[str, Any],
) -> str:
    words = text.split()
    if len(words) > max_words:
        return " ".join(words[:max_words])
    if len(words) >= min_words:
        return text

    additions: List[str] = []
    remaining_gaps = resume_json.get("remaining_gaps", [])
    if remaining_gaps:
        additions.append(
            "I am careful to present my fit honestly, including where the role asks for depth that would need to be demonstrated in discussion rather than overstated in writing."
        )
    if candidate_profile.get("experience_years"):
        additions.append(
            f"My {int(candidate_profile.get('experience_years', 0) or 0)} years of experience have reinforced a disciplined approach to delivery, collaboration, and quality ownership."
        )
    if metadata.get("source"):
        additions.append(
            f"I appreciate the chance to be considered through {metadata.get('source')} and would be glad to provide any additional details needed for review."
        )

    enriched = text
    for sentence in additions:
        if len(enriched.split()) >= min_words:
            break
        enriched = enriched.replace("\n\nThank you", f" {sentence}\n\nThank you")
    return enriched


def _validate_cover_letter(text: str, candidate_profile: Dict[str, Any]) -> None:
    word_count = len(text.split())
    if word_count < MIN_WORDS or word_count > MAX_WORDS:
        raise CoverLetterGenerationError(
            f"Generated cover letter does not satisfy the {MIN_WORDS}-{MAX_WORDS} word constraint"
        )
    if "%" in text:
        raise CoverLetterGenerationError("Generated cover letter contains unsupported numeric claims")
    candidate_name = str(candidate_profile.get("name", "")).strip()
    if candidate_name and candidate_name not in text:
        raise CoverLetterGenerationError("Generated cover letter is missing the candidate name signature")


def _write_docx(output_path: Path, text: str, metadata: Dict[str, Any]) -> None:
    document = Document()
    document.add_heading(f"Cover Letter - {metadata.get('role_title', 'Role')}", level=1)
    for paragraph in text.split("\n\n"):
        document.add_paragraph(paragraph.strip())
    document.save(output_path)


def _write_pdf(output_path: Path, text: str, metadata: Dict[str, Any]) -> None:
    pdf = canvas.Canvas(str(output_path), pagesize=letter)
    width, height = letter
    top = height - 0.75 * inch
    left = 0.75 * inch
    y = top
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(left, y, f"Cover Letter - {metadata.get('role_title', 'Role')}")
    y -= 20
    pdf.setFont("Helvetica", 10)
    for paragraph in text.split("\n\n"):
        for line in _wrap_text(paragraph, 95):
            if y < 0.75 * inch:
                pdf.showPage()
                y = top
                pdf.setFont("Helvetica", 10)
            pdf.drawString(left, y, line)
            y -= 14
        y -= 8
    pdf.save()


def _wrap_text(text: str, width: int) -> List[str]:
    words = text.split()
    if not words:
        return [""]
    lines = []
    current = words[0]
    for word in words[1:]:
        if len(current) + 1 + len(word) <= width:
            current = f"{current} {word}"
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def _require_dict(value: Any, field_name: str) -> Dict[str, Any]:
    if not isinstance(value, dict):
        raise CoverLetterGenerationError(f"Field '{field_name}' must be an object")
    return value


def _require_text(value: Any, field_name: str) -> str:
    text = str(value or "").strip()
    if not text:
        raise CoverLetterGenerationError(f"Missing required field: {field_name}")
    return text